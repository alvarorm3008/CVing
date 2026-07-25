from io import BytesIO

from docx import Document
from fastapi import HTTPException


def _collect_hyperlink_urls(document: Document) -> list[str]:
    urls: list[str] = []
    seen: set[str] = set()
    try:
        for rel in document.part.rels.values():
            if "hyperlink" not in (rel.reltype or ""):
                continue
            target = (rel.target_ref or "").strip()
            if not target:
                continue
            key = target.lower().rstrip("/")
            if key in seen:
                continue
            seen.add(key)
            urls.append(target)
    except Exception:
        pass
    return urls


def _iter_paragraph_texts(paragraphs) -> list[str]:
    parts: list[str] = []
    for paragraph in paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    return parts


def extract_text_from_docx(file_bytes: bytes) -> str:
    if not file_bytes:
        raise HTTPException(status_code=400, detail="The uploaded DOCX file is empty.")

    try:
        document = Document(BytesIO(file_bytes))
    except Exception as exc:
        raise HTTPException(status_code=400, detail="Could not read the DOCX file.") from exc

    parts: list[str] = []

    # Headers / footers often hold contact links
    for section in document.sections:
        parts.extend(_iter_paragraph_texts(section.header.paragraphs))
        parts.extend(_iter_paragraph_texts(section.footer.paragraphs))

    parts.extend(_iter_paragraph_texts(document.paragraphs))

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    body = "\n".join(parts).strip()
    body_lower = body.lower()
    extras: list[str] = []
    for url in _collect_hyperlink_urls(document):
        needle = url
        if url.lower().startswith("mailto:"):
            needle = url.split(":", 1)[1]
        if needle.lower() not in body_lower and url.lower() not in body_lower:
            extras.append(url)

    if extras:
        body = f"{body}\n\n" + "\n".join(extras) if body else "\n".join(extras)

    extracted = body.strip()
    if not extracted:
        raise HTTPException(
            status_code=400,
            detail="Could not extract text from the DOCX. Make sure it contains readable content.",
        )

    return extracted
